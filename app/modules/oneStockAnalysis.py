from typing import List

import pandas as pd
import math
import time

from docx import Document
from docx.shared import Inches, Pt
from quickchart import QuickChart
import tushare as ts

# 导入本地库
from . import config
from ..logging import logger
from .models.FamaFrench import FamaFrench

# 初始化pro接口
pro = ts.pro_api(config.tushareToken)

# 储存所有信息
DATA = {}


# 撰写内容
def write(document: Document, path: str) -> Document():
    logger.info('[个股分析] 模块开始载入 ...')

    # 添加标题
    document.add_heading('个股分析', level=1)

    # 从指定指数中抽一只股票
    stock = pickStock('000300.SH')
    stockName = getStockName(stock)
    DATA['stockName'] = stockName
    logger.info(f'[个股分析] 抽选股票名称 {stockName} 代码 {stock} 。')

    # Fama French 归因
    document.add_heading('Fama-French 三因子归因分析', level=2)
    logger.info(f'[个股分析] 开始 Fama French 三因子归因 ...')
    t = time.time()
    ff = FamaFrench(stock)
    ff.pull()
    logger.info(f'[个股分析] 成功抓取近一年行情数据 ...')
    ff.build()
    ff.regression()
    logger.info('[个股分析] Fama French 归因完成，用时 %s 秒！' % ('%.2f' % (time.time() - t)))
    document.add_paragraph(f'{stockName} ({stock[:6]}) 的归因结果：')
    document.add_paragraph(f'市场风险因子：{"%.4f" % ff.weights[0]}')
    document.add_paragraph(f'规模风险因子：{"%.4f" % ff.weights[1]}')
    document.add_paragraph(f'账面市值比风险因子：{"%.4f" % ff.weights[2]}')

    famaChart(ff, path)
    document.add_picture(DATA['ffChart'], width=Inches(6))


# 验证今日是否添加
def dayCheck() -> bool:
    return True


# 抽选个股
def pickStock(index: str, blacklist: List[str] = None):
    """
    随机抽选一只指数成分股
    :return: 股票代码（含.TS）或None（指定指数不存在或无成分股）
    """

    if blacklist is None:
        blacklist = []

    # 获取指数成分股
    # 先检验一下这个指数是否存在以及最近更新日期
    df = pro.index_weight(**{
        "index_code": index,
        "trade_date": "",
        "start_date": "",
        "end_date": "",
        "limit": 1,
        "offset": ""
    }, fields=[
        "index_code",
        "con_code",
        "trade_date",
        "weight"
    ])

    # 检查是否没有返回数据
    if df.index.size == 0:
        logger.error('[个股分析] 指定指数代码%s不存在或无行情信息，请检查！' % index)
        return None

    # 获取最新更新日期
    date = df['trade_date'][0]

    # 获取所有指数和权重
    df = pro.index_weight(**{
        "index_code": index,
        "trade_date": date,
        "start_date": "",
        "end_date": "",
        "limit": "",
        "offset": ""
    }, fields=[
        "index_code",
        "con_code",
        "trade_date",
        "weight"
    ])

    sample = df.sample()

    while df.index.size > 0:
        if sample['con_code'] in blacklist:
            df.drop(df.index(df['con_code'] == sample['con_code']), axis='index', inplace=True)
            sample = df.sample()
        else:
            break

    del df
    return sample['con_code'].iloc[0]


def getStockName(code: str):
    # 拉取数据
    df = pro.stock_basic(**{
        "ts_code": code,
        "name": "",
        "exchange": "",
        "market": "",
        "is_hs": "",
        "list_status": "",
        "limit": "",
        "offset": ""
    }, fields=[
        "ts_code",
        "symbol",
        "name",
        "area",
        "industry",
        "market",
        "list_date"
    ])

    if df.index.size == 0:
        return None

    return df['name'][0]


def famaChart(ff, path: str):

    qc = QuickChart()
    qc.width = 750
    qc.height = 300
    qc.version = '2'

    # Config can be set as a string or as a nested dict
    qc.config = """{
      type: 'line',
      data: {
        labels: %s,
        datasets: [
          {
            label: '%s',
            borderColor: 'rgb(0, 92, 230)',
            pointRadius: 0,
            fill: false,
            data: %s
          },
          {
            label: '市场因子',
            borderColor: 'rgb(0, 153, 77)',
            pointRadius: 0,
            fill: false,
            data: %s
          },
          {
            label: '规模因子',
            borderColor: 'rgb(255, 128, 0)',
            pointRadius: 0,
            fill: false,
            data: %s
          },
          {
            label: 'PB因子',
            borderColor: 'rgb(153, 51, 255)',
            pointRadius: 0,
            fill: false,
            data: %s
          },
          {
            label: '无风险收益率',
            borderColor: 'rgb(191, 191, 191)',
            pointRadius: 0,
            fill: false,
            data: %s
          }
        ],
      },
      options: {
        title: {
          display: false
        },
      },
    }""" % (ff.data['trade_date'].values.tolist(), DATA['stockName'], ff.data['r'].values.tolist(),
            ff.data['mkt_rf'].values.tolist(), ff.data['smb'].values.tolist(), ff.data['hml'].values.tolist(),
            ff.data['rf'].values.tolist())

    # 输出一次图片链接地址，可以在浏览器打开
    logger.debug(qc.get_url())
    logger.info('[个股分析] Fama French 因子收益图生成中 ...')
    t = time.time()

    # Or write the chart to a file
    qc.to_file(path + 'ffChart.png')
    logger.info('[个股分析] Fama French 因子收益图已生成，用时 %s 秒！' % ('%.2f' % (time.time() - t)))

    # 在data中确认图片地址
    DATA['ffChart'] = path + 'ffChart.png'