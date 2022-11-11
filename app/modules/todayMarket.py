import pandas as pd
import math
import datetime
import calendar
import time

from docx import Document
from docx.shared import Inches, Pt
from quickchart import QuickChart
import tushare as ts

# 导入本地库
from . import config
from ..logging import logger
from ..style import tableRowHeightSet, tableCenterSet

# 初始化pro接口
pro = ts.pro_api(config.tushareToken)


# 撰写内容
def write(document: Document, path: str) -> Document():
    logger.info('[今日市场] 模块开始载入 ...')

    # 添加标题
    document.add_heading('今日市场', level=1)

    # 执行分析
    process(path)

    # region 添加股票涨跌数量比较
    if 'marketImage' in data:
        logger.info('[今日市场] 写入股票涨跌 ...')
        document.add_heading('股票涨跌', level=2)
        document.add_picture(data['marketImage'], width=Inches(6))
        paragraph = document.add_paragraph('今日，共有 ')
        paragraph.add_run(str(data['stocks']['increase']) + ' 只股票上涨， ')
        paragraph.add_run(str(data['stocks']['decrease']) + ' 只股票下跌， ')
        paragraph.add_run(str(data['stocks']['decrease']) + ' 只股票持平或停牌。')
        paragraph.paragraph_format.line_spacing = 1.5
    # endregion

    # region 添加市场风格统计
    if 'style' in data:
        logger.info('[今日市场] 写入指数涨跌 ...')
        document.add_heading('指数涨跌统计', level=2)
        table = document.add_table(rows=1, cols=4, style='Colorful List Accent 4')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指数名称'
        hdr_cells[1].text = '指数收盘价'
        hdr_cells[2].text = '涨跌幅'
        hdr_cells[3].text = '成交额（亿元）'
        for r in data['style']:
            row_cells = table.add_row().cells
            row_cells[0].text = r['name']
            row_cells[1].text = format(float(r['price']), ',')
            row_cells[2].text = r['change']
            row_cells[3].text = format(float(r['amount']), ',')

        # 调整表格行高为400
        tableRowHeightSet(table, 500)
        # 设置表格水平居中、垂直居中
        tableCenterSet(table)

    # endregion

    logger.info('[今日市场] 模块运行完成！')

    return document


# 验证今日是否添加
def dayCheck() -> bool:
    logger.info('[今日市场] 检测日期有效性 ...')

    # 获取今日日期
    currentDate = datetime.date.today()
    currentDay = calendar.weekday(currentDate.year, currentDate.month, currentDate.day)

    # 判断是否是工作日，如果是周末就返回False
    if currentDay > 5:
        logger.info('[今日市场] 当日无效，跳过中 ...')
        return False
    else:
        logger.info('[今日市场] 当日有效，加载中 ...')
        return True


# 这个是所有数据存入的地方
data = {}


def process(path: str):
    # 市场涨跌统计
    changeStatistic(path)

    # 市场风格统计
    styleStatistic()


def changeStatistic(path: str):
    """
    有关市场整体涨跌情况的统计，生成一张涨跌分布图。
    """

    logger.info('[今日市场] 开始统计整个市场个股当日涨跌...')
    t = time.time()

    # 拉取数据
    df = pro.daily(**{
        "ts_code": "",
        "trade_date": 20221111,
        "start_date": "",
        "end_date": "",
        "offset": "",
        "limit": ""
    }, fields=[
        "ts_code",
        "pct_chg",
        "vol",
        "amount"
    ])

    logger.debug(df)
    logger.info('[今日市场] 成功加载今日个股数据，用时 %s 秒，载入个股数量 %s 只！' % ('%.2f' % (time.time() - t), df.iloc[:, 0].size))

    # 保留0 3 6开头的股票
    df = df[df.ts_code.str.contains("[0,3,6][0-9]{5}", regex=True)]
    df['ts_code'] = df['ts_code'].str[0:6]

    # 获取涨跌幅情况
    df['rank'] = df["pct_chg"].map(lambda x: min(max(math.floor(x), -11), 11) + 11)

    logger.debug(df)
    logger.info('[今日市场] 完成个股数据筛选，保留个股数量 %s 只！' % df.iloc[:, 0].size)

    # 统计涨跌幅分布
    count = list(df.groupby('rank')['pct_chg'].count())
    green = count[:11] + [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]
    red = [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0] + count[11:]

    # 统计涨跌个数
    increaseNumber = df[df['pct_chg'] > 0].index.size
    decreaseNumber = df[df['pct_chg'] < 0].index.size
    noChangeNumber = df[df['pct_chg'] == 0].index.size
    data['stocks'] = {
        'increase': increaseNumber,
        'decrease': decreaseNumber,
        'noChange': noChangeNumber
    }

    # 开始作图
    qc = QuickChart()
    qc.width = 750
    qc.height = 300
    qc.version = '2'

    # Config can be set as a string or as a nested dict
    qc.config = """{
      type: 'bar',
      data: {
        labels: ['<-10', '10', '-9', '-8', '-7', '-6', '-5', '-4', '-3', '-2', '-1', '0', '1', 
            '2', '3', '4', '5', '6', '7', '8', '9', '-10', '>10'],
        datasets: [
          {
            label: 'Dataset 1',
            backgroundColor: 'rgb(2, 112, 48)',
            data: %s
          },
          {
            label: 'Dataset 2',
            backgroundColor: 'rgb(173, 10, 29)',
            data: %s
          }
        ],
      },
      options: {
        title: {
          display: false
        },
        legend: {
          display: false
        },
        scales: {
          xAxes: [
            {
              stacked: true
            },
          ],
          yAxes: [
            {
              ticks: {
                max: %s,
                stepSize: 100
              },
              stacked: true
            },
          ],
        },
        plugins: {
          datalabels: {
            anchor: 'end',
            align: 'top',
            formatter: (value) => {
              if (value > 0) 
                return value
              else 
                return ''
            }
          },
        },
      },
    }""" % (green, red, (math.floor(max(count) / 200) + 2) * 200)

    # 输出一次图片链接地址，可以在浏览器打开
    logger.debug(qc.get_url())
    logger.info('[今日市场] 市场个股涨跌统计图生成中 ...')
    t = time.time()

    # Get the image as a variable...
    qc.to_file(path + 'marketImage.png')

    logger.info('[今日市场] 市场个股涨跌统计图已生成，用时 %s 秒！' % ('%.2f' % (time.time() - t)))

    # 在data中确认图片地址
    data['marketImage'] = path + 'marketImage.png'


def styleStatistic():
    """
    市场风格的统计分析
    """

    logger.info('[今日市场] 开始统计风格指数当日行情...')
    t = time.time()

    # 风格指数的数据存进这个array，之后要根据涨跌幅排序
    results = []

    # 定义需要抓取的大盘指数
    indexCode = {
        '000001.SH': '上证综指',
        '000016.SH': '上证50',
        '000300.SH': '沪深300',
        '000905.SH': '中证500',
        '000852.SH': '中证1000',
        '399006.SZ': '创业板指',
        '000688.SH': '科创50'
    }

    for i in indexCode.keys():
        # 拉取数据
        result = pro.index_daily(**{
            "ts_code": i,
            "trade_date": 20221111,
            "start_date": "",
            "end_date": "",
            "limit": "",
            "offset": ""
        }, fields=[
            "ts_code",
            "trade_date",
            "close",
            "open",
            "high",
            "low",
            "pre_close",
            "change",
            "pct_chg",
            "vol",
            "amount"
        ])

        # 将数据添加进去
        results.append(result)

    # 整合所有dataframe
    df = pd.concat(results)

    logger.info('[今日市场] 完成指数行情加载，用时 %s 秒，共统计 %s 只指数！' % ('%.2f' % (time.time() - t), df.iloc[:, 0].size))

    # 成交额和成交量按亿为单位统计
    df['vol'] = df['vol'].map(lambda x: x * 100 / 100000000)
    df['amount'] = df['amount'].map(lambda x: x / 100000)

    # 根据涨跌幅降序排列
    # 主要是看看每天是什么风格涨的高
    df = df.sort_values(by=['pct_chg'], ascending=False)
    df = df.reset_index(drop=True)

    # 遍历df开始往data里写东西
    data['style'] = []
    for index, row in df.iterrows():
        data['style'].append({
            'code': row['ts_code'],
            'name': indexCode[row['ts_code']],
            'price': '%.2f' % row['close'],
            'change': '%.2f' % row['pct_chg'] + '%',
            'amount': '%.2f' % row['amount']
        })
