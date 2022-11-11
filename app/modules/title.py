from docx import Document


# 撰写内容
def write(document: Document, path: str) -> Document():
    document.add_heading('QC Daily', 0)
    document.add_paragraph('Quant Cat knows everything!')

    return document


# 验证今日是否添加
def dayCheck() -> bool:
    return True
