from docx import Document


# 撰写内容
def write(document: Document, path: str) -> Document():
    document.add_heading('声明', level=1)
    paragraph = document.add_paragraph(
        '本报告内容由自动化程序制作而成，基础数据来源于', style='List Bullet'
    )
    paragraph.add_run('公开数据').underline = True
    paragraph.add_run('和')
    paragraph.add_run('授权数据').underline = True
    paragraph.add_run('，本文作者及程序作者均不对数据准确性和时效性做出任何保证。')
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph = document.add_paragraph(
        '本程序仅用于学习及研究使用，不构成任何投资建议，因本文内容做出的投资决策导致的损失，本文作者及程序作者概不负责。',
        style='List Bullet'
    )
    paragraph.paragraph_format.line_spacing = 1.5
    return document


# 验证今日是否添加
def dayCheck() -> bool:
    return True
