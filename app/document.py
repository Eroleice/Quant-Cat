from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from . import modules


# 新建文档
def new() -> Document():
    # 新建文档
    document = Document()

    # 设置字体
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(12)
    font.name = u'Arial'  # 设置西文字体
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线')  # 设置中文字体

    return document


# 读取各个module
def addModule(document: Document, path: str) -> Document():
    for m in modules.__all__:
        if m.dayCheck():
            m.write(document, path)
    return document
